
import streamlit as st
from PIL import Image
from docx import Document
from datetime import datetime
import pytesseract
import re
import io

# Configura o Tesseract no Windows, se necessário
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

procedimentos = [
    "Alta Compartilhada",
    "Atendimento pós operatório cirurgia bariátrica",
    "Atendimento psicológico acompanhante/familiares",
    "Atendimento psicológico beira leito ao paciente",
    "Atendimento psicologico emergencial",
    "Auxílio no comunicado de diagnóstico e prognóstico",
    "Auxílio no comunicado de óbito aos familiares",
    "Avaliação psicológica bariátrica",
    "Avaliação psicológica para visita estendida",
    "Discussão de caso e planejamento da agenda bariátrica",
    "Entrega de relatório psicológico bariátrica",
    "Entrevista suporte afetivo familiar bariátrica",
    "Encaminhamento para Acompanhamento Psicológico ao paciente",
    "Encaminhamento para Acompanhamento Psicológico aos familiares",
    "Orientação as normas da instituição - HRCLMT",
    "Pocesso seletivo- entrevista RH",
    "Visita multidisciplinar/multiprofissional (UTI e clínicas)",
    "Prontuário Afetivo",
    "Acompanhante terapêutico"
]

def traduzir_numeros(indices):
    return [procedimentos[i - 1] for i in indices if 0 < i <= len(procedimentos)]

def extrair_dados(imagem):
    texto = pytesseract.image_to_string(imagem, lang='por')
    padrao = re.compile(r"([A-ZÁÉÍÓÚÃÕÇ\s]+)\s+(\d+(?:,\s*\d+)*)")
    dados = []
    for linha in texto.split("\n"):
        match = padrao.search(linha.strip())
        if match:
            nome = match.group(1).strip()
            numeros = [int(n.strip()) for n in match.group(2).split(',') if n.strip().isdigit()]
            dados.append((nome.title(), traduzir_numeros(numeros)))
    return dados

def gerar_documento(responsavel, crp, horario, censo, atendimentos):
    doc = Document()
    data = datetime.now().strftime("%d/%m/%Y")
    doc.add_paragraph(f"PASSAGEM DE PLANTÃO DE PSICOLOGIA\n\n   Data: {data} \n   Horário: {horario}\n\n   Responsável pelo Plantão: {responsavel}             CRP {crp}\n\n")

    doc.add_paragraph("Censo hospitalar\n")
    doc.add_paragraph(f"Clínica cirúrgica: {censo['cirurgica']}             UTI adulto: {censo['uti_adulto']}\n")
    doc.add_paragraph(f"Clínica médica: {censo['medica']}                UTI coronária: {censo['uti_coronaria']}\n")
    doc.add_paragraph(f"Clínica pediátrica: {censo['pediatrica']}           UTI pediátrica: {censo['uti_pediatrica']}\n\n")

    doc.add_paragraph("Atendimentos por setor:\n")
    doc.add_paragraph("UTI CORONÁRIA\n")
    for nome, proc in atendimentos:
        doc.add_paragraph(f"{nome} — Procedimentos: {', '.join(proc)}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- APP ---
st.title("📝 Geração Automática de Ata Psicológica")

imagem = st.file_uploader("Envie a foto do censo com os procedimentos anotados:", type=["png", "jpg", "jpeg"])

with st.form("formulario"):
    responsavel = st.text_input("Responsável pelo Plantão")
    crp = st.text_input("CRP")
    horario = st.text_input("Horário do plantão (ex: 07hrs às 13hrs)")

    st.markdown("### Censo Hospitalar")
    censo = {
        'cirurgica': st.text_input("Clínica cirúrgica"),
        'uti_adulto': st.text_input("UTI adulto"),
        'medica': st.text_input("Clínica médica"),
        'uti_coronaria': st.text_input("UTI coronária"),
        'pediatrica': st.text_input("Clínica pediátrica"),
        'uti_pediatrica': st.text_input("UTI pediátrica")
    }

    enviado = st.form_submit_button("Gerar Ata")

if imagem and enviado:
    img = Image.open(imagem)
    dados = extrair_dados(img)

    if not dados:
        st.error("Nenhum dado foi reconhecido na imagem. Verifique a nitidez e a escrita dos números.")
    else:
        docx_file = gerar_documento(responsavel, crp, horario, censo, dados)
        st.success("Ata gerada com sucesso!")
        st.download_button("📥 Baixar Ata em .docx", data=docx_file, file_name="Ata_Psicologia.docx")
